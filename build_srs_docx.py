import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_header(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Set Tab Stop at 16cm (Right margin: 21cm - 3cm left - 2cm right = 16cm)
        tab_stops = p.paragraph_format.tab_stops
        try:
            tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            pass # Fallback if tab_stops addition fails
            
        run = p.add_run("DỰ ÁN KHU DU LỊCH ĐẠI NAM\tĐẶC TẢ YÊU CẦU NGHIỆP VỤ (SRS)")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(31, 73, 125) # 1F497D
        
        # Add a thick bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # Thicker border (12 eighths of a point)
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)

def create_footer(doc):
    # Add footer with page number
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Trang ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
        # Add PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    
    # Set table to span 100% of the window width
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000') # 5000 means 100%
    tblPr.append(tblW)
    
    # Fill empty STT header if 3 columns
    if len(headers) == 3 and (headers[0].strip() == '' or headers[0].strip().upper() == 'STT'):
        headers[0] = 'STT'
                
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Background color for header
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D') # Standard Word Blue
        cell._tc.get_or_add_tcPr().append(shd)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternate row colors
            if ri % 2 == 1:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1') # Standard Word Light Blue
                cell._tc.get_or_add_tcPr().append(shd)
            
            p = cell.paragraphs[0]
            # Replace markdown bold with actual bold
            parts = re.split(r'(\*\*.*?\*\*)', val)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            
            # Auto-align STT column (column 0)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_srs_docx():
    doc = Document()
    
    # 1. Page setup (A4)
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
    
    # Header and Footer
    create_header(doc)
    create_footer(doc)
    
    # 2. Styles
    styles = doc.styles
    def set_style(style_name, font_name, font_size, bold=False, color=None, space_before=0, space_after=6):
        try:
            s = styles[style_name]
        except KeyError:
            return
        s.font.name = font_name
        s.font.size = Pt(font_size)
        s.font.bold = bold
        if color:
            s.font.color.rgb = RGBColor(*color)
        s.paragraph_format.space_before = Pt(space_before)
        s.paragraph_format.space_after  = Pt(space_after)

    set_style('Normal',    'Times New Roman', 13, space_before=6)
    set_style('Heading 1', 'Times New Roman', 18, bold=True, color=(0, 70, 127), space_before=24, space_after=12)
    set_style('Heading 2', 'Times New Roman', 15, bold=True, color=(31, 73, 125), space_before=18, space_after=8)
    set_style('Heading 3', 'Times New Roman', 14, bold=True, color=(17, 85, 204), space_before=14, space_after=6)
    set_style('Heading 4', 'Times New Roman', 13, bold=True, color=(0, 0, 0), space_before=12, space_after=6)
    
    # 3. Cover Page
    def empty_lines(count):
        for _ in range(count):
            doc.add_paragraph()
            
    empty_lines(3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CÔNG TY CỔ PHẦN ĐẠI NAM')
    run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TRUNG TÂM PHÁT TRIỂN PHẦN MỀM')
    run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.bold = True
    
    empty_lines(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM (SRS)')
    run.font.name = 'Times New Roman'; run.font.size = Pt(22); run.bold = True
    run.font.color.rgb = RGBColor(0, 70, 127)
    
    empty_lines(1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Phân hệ: HỆ THỐNG & DANH MỤC')
    run.font.name = 'Times New Roman'; run.font.size = Pt(18); run.bold = True
    
    empty_lines(5)
    
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    
    def set_cell(row_idx, col_idx, text, bold=False):
        c = tbl.rows[row_idx].cells[col_idx]
        p = c.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = bold
        
    set_cell(0, 0, 'Mã dự án: ', True); set_cell(0, 1, 'DN01')
    set_cell(1, 0, 'Mã tài liệu: ', True); set_cell(1, 1, 'DN01_SRS_HeThong_DanhMuc_v1.3')
    set_cell(2, 0, 'Nhóm thực hiện: ', True); set_cell(2, 1, 'Nhóm Phân tích Nghiệp vụ (BA)')
    
    empty_lines(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tháng 04/2026')
    run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.italic = True
    
    doc.add_page_break()
    
    # 4. Parse Markdown Document
    with open('docs/Sprint1/07_SRS_DanhMuc_Merged.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        raw_line = line.strip()
        if not raw_line:
            if in_table:
                # Flush table
                add_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            continue
            
        # Skip the original markdown cover lines
        if raw_line.startswith('# Khu Du Lịch Đại Nam') or \
           raw_line.startswith('# Đặc Tả Yêu Cầu Phần Mềm') or \
           raw_line.startswith('# Mã dự án:') or \
           raw_line.startswith('# Mã tài liệu:') or \
           raw_line == 'Hồ Chí Minh, Tháng 04/2026':
            continue
            
        # Table Parsing
        if raw_line.startswith('|') and raw_line.endswith('|'):
            cells = [cell.strip() for cell in raw_line.strip('|').split('|')]
            if set(raw_line) == {'|', '-'}:
                pass # Skip the separator row
            elif not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
            
        if in_table:
            # Table ended but no empty line
            add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
            
        # Headings
        if raw_line == '---':
            continue
            
        if raw_line.startswith('# '):
            p = doc.add_heading(raw_line[2:], level=1)
            p.runs[0].font.name = 'Times New Roman'
            # Add page break before H1 unless it's the very first content
            if len(doc.paragraphs) > 20:
                p.insert_paragraph_before('').add_run().add_break()
        elif raw_line.startswith('## '):
            p = doc.add_heading(raw_line[3:], level=2)
            p.runs[0].font.name = 'Times New Roman'
        elif raw_line.startswith('### '):
            p = doc.add_heading(raw_line[4:], level=3)
            p.runs[0].font.name = 'Times New Roman'
        elif raw_line.startswith('#### '):
            p = doc.add_heading(raw_line[5:], level=4)
            p.runs[0].font.name = 'Times New Roman'
            
        # Lists
        elif raw_line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(raw_line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            
        # Placeholders
        elif raw_line.startswith('[Placeholder:'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(raw_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add a subtle box effect
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for border in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '4')
                b.set(qn('w:color'), 'CCCCCC')
                pBdr.append(b)
            pPr.append(pBdr)
            
        # Normal Paragraph
        else:
            # Handle bold text
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', raw_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)

    out_path = 'docs/Sprint1/07_SRS_DanhMuc_Merged_v7.docx'
    try:
        doc.save(out_path)
        print(f'Successfully generated Word document: {out_path}')
    except PermissionError:
        out_path = 'docs/Sprint1/07_SRS_DanhMuc_Merged_v7_alt.docx'
        doc.save(out_path)
        print(f'Saved as alternate: {out_path}')

if __name__ == '__main__':
    build_srs_docx()
